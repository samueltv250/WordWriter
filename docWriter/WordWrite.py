import docx
import os
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from extract_Alcance_PH import *
from numero_letras import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
#
#doc = docx.Document()
#h1 = doc.add_heading('Test Doc',0)
#parag = doc.add_paragraph("Hello!")
#italicparagraph = doc.add_paragraph()
#h1.add_run("test123, ")
#parag.add_run("This word document was created using, ")
#parag.add_run("Python").bold=True
#doc.add_paragraph("")
#
#doc.add_paragraph("")
#doc.add_heading('Heading Level 2',2)
#
#italicparagraph.add_run("This line is in italics!").italic=True
#parag.add_run("Python").bold=True
#doc.save("test.docx")
#os.system("start test.docx")

import itertools

dict_letters = {}
listofdatt , tipomod , dicdeff, dicTOT= getVals()

for p in dicdeff:
    x = p.split()
    if x[0] not in dict_letters:
        dict_letters[x[0]] = []
    
    if len(x) != 1:
        dict_letters[x[0]].append(x[1])
    
list_of_phase = []
countofPairs = 0
for i in tipomod:
    for p in tipomod[i]:
        list_of_phase.append(p)
        
#list_of_phase = sorted(list_of_phase)
totcount_phase = len(list_of_phase)


dictOfLists = {}
for p in list_of_phase:
    if p[0] not in dictOfLists:
        dictOfLists[p[0]] = [int(p[2:])]
    else:
        dictOfLists[p[0]].append(int(p[2:]))

for i in dictOfLists:
    dictOfLists[i] = sorted(dictOfLists[i])

def ranges(i):
    for a, b in itertools.groupby(enumerate(i), lambda pair: pair[1] - pair[0]):
        b = list(b)
        yield b[0][1], b[-1][1]



def rangesList(dictOfLists,countofPairs):
    for i in dictOfLists:
        dictOfLists[i] =  list(ranges(dictOfLists[i]))
        countofPairs += len(dictOfLists[i])
    return dictOfLists, countofPairs

        
        






list_of_phase, countofPairs = rangesList(dictOfLists,countofPairs)

print(list_of_phase)



print(totcount_phase)





document = docx.Document()
document.styles['Normal'].font.name = 'Arial'

style = document.styles.add_style('bca', WD_STYLE_TYPE.PARAGRAPH)
style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
font = style.font
font.bold = True
font.name = 'Arial'

style = document.styles.add_style('bcal', WD_STYLE_TYPE.PARAGRAPH)
style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
font = style.font
font.underline = True
font.bold = True
font.name = 'Arial'

style = document.styles.add_style('ba', WD_STYLE_TYPE.PARAGRAPH)
style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
font = style.font
font.bold = True
font.name = 'Arial'

style = document.styles.add_style('bal', WD_STYLE_TYPE.PARAGRAPH)
#style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
font = style.font
font.underline = True
font.bold = True
font.name = 'Arial'

def add_Title(document,listofdatt):
    paragrap = document.add_paragraph("REGLAMENTO DE COPROPIEDAD",style='bca')
    paragrap = document.add_paragraph("DEL "+str(listofdatt[0][0]),style='bca')
    paragrap = document.add_paragraph("CAPITULO I",style='bcal')
    paragrap = document.add_paragraph(str())
    paragrap = document.add_paragraph(str("DISPOSICIONES GENERALES"),style='bcal')
    
    
def add_Art_1(document,listofdatt):
    paragraph = document.add_paragraph("")
    run = paragraph.add_run("ARTICULO 1: ")
    run.underline = True
    run.bold = True
    paragraph.add_run("Quedan sometidas al Régimen de Propiedad Horizontal con arreglo a las Disposiciones Legales previstas en la Ley doscientos ochenta y cuatro (284) del catorce (14) de febrero de dos mil veintidós (2022) (en adelante La Ley) y demás disposiciones pertinentes, la Finca inscrita al folio real  número ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][3])))
    run.bold = True
    paragraph.add_run(" con código de ubicación ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][4])))
    run.bold = True
    paragraph.add_run(", ubicada en el ")
    
    run = paragraph.add_run(str(listofdatt[0][5]))
    run.bold = True
    paragraph.add_run(", inscrita la sección de la Propiedad, del Registro Público, Provincia de Panamá, así como las mejoras construidas sobre dicha finca; en adelante, denominada ")
    run = paragraph.add_run(str(listofdatt[0][0])+".")
    run.bold = True

def add_Art_2(document,listofdatt):
    paragraph = document.add_paragraph("")
    run = paragraph.add_run("ARTICULO 2: ")
    run.underline = True
    run.bold = True
    paragraph.add_run("EL ")
    run = paragraph.add_run(str(listofdatt[0][0]))
    run.bold = True
    paragraph.add_run(" está ubicado en ")
    run = paragraph.add_run(str(listofdatt[0][16]))
    run.bold = True
    paragraph.add_run(", y la conforma un lote de terreno con una superficie de ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][6])))
    run.bold = True
    paragraph.add_run(" y un valor de ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][7])))
    run.bold = True
    paragraph.add_run(".")
    
    
    
    document.add_paragraph("")
    paragraph = document.add_paragraph("")
    paragraph.add_run("Sobre ella se ha construido un conjunto de unidades inmobiliarias residenciales, el cual consiste en un desarrollo urbanístico residencial, por lo cual será de ")
    run = paragraph.add_run("USO RESIDENCIAL")
    run.bold = True
    paragraph.add_run(", el conjunto de unidades inmobiliarias residenciales, ha sido construido con estructura de hormigón, paredes de concreto armado repelladas en ambas caras, piso de baldosa, ventana con marco de aluminio y vidrio corredizo, techo de fibrocemento o láminas de acero galvanizado (según el modelo), el cual será denominado en su conjunto como ")
    run = paragraph.add_run(str(listofdatt[0][0]))
    run.bold = True
    run = paragraph.add_run(".")
    run.bold = True
    
    
    document.add_paragraph("")
    paragraph = document.add_paragraph("La construcción del proyecto ocupa una superficie de ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][6])))
    run.bold = True
    paragraph.add_run(",  de los cuales, ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][8])))
    run.bold = True
    paragraph.add_run(" corresponden al área de construcción de las unidades inmobiliarias, ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][9])))
    run.bold = True
    paragraph.add_run(" corresponden a las áreas comunes entrada vehicular y peatonal, garita de entrada, calles y vía de acceso, aceras, área de parque y áreas verdes")
    
    if str(listofdatt[0][10]) != "na":
        run = paragraph.add_run(" y "+get_numbandword(str(listofdatt[0][10])))
        run.bold = True
        paragraph.add_run(" serán reserva del promotor, quedando en la finca propiedad ")
        run = paragraph.add_run(get_numbandword(str(listofdatt[0][3])))
        run.bold = True
        paragraph.add_run(", con código de ubicación ")
        run = paragraph.add_run(get_numbandword(str(listofdatt[0][4])))
        run.bold = True
    
    paragraph.add_run(".")
    
    
    
def add_Desc_Gen(document,listofdatt,dicTOT):


    paragraph = document.add_paragraph("DESCRIPCION GENERAL DEL P.H. MONTEMADERO V.",style='ba')
    paragraph = document.add_paragraph("")
    run = paragraph.add_run(str(listofdatt[0][15]))
    document.add_paragraph("")
    
    
    
    
    paragraph = document.add_paragraph("")
    run = paragraph.add_run("El ")
    run = paragraph.add_run(str(listofdatt[0][0])+",")
    run.bold = True
    run = paragraph.add_run(" estará compuesto por ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][2])))
    run = paragraph.add_run(" unidades inmobiliarias en total, las cuales serán identificadas de la ")
    run = paragraph.add_run("E-001 a la E-"+str(listofdatt[0][2]) + ".")
        
    paragraph = document.add_paragraph("")
    
    paragraph = document.add_paragraph("")
    run = paragraph.add_run("El ")
    run = paragraph.add_run(str(listofdatt[0][0])+",")
    run.bold = True
    run = paragraph.add_run(" estará compuesto por ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][2])))
    run = paragraph.add_run(" unidades inmobiliarias, desglosadas así:")
    for i in dicTOT:
        run = paragraph.add_run(" "+get_numbandword(str(dicTOT[i])))
        run = paragraph.add_run(" unidades modelo ")
        run = paragraph.add_run(str(i)+",")

    run = paragraph.add_run(" las cuales se describen a continuación:")

    
    
def add_Desc_Rep(document,listofdatt,dicdeff):
    for i in dicdeff:
        paragraph = document.add_paragraph("UNIDAD INMOBILIARIA ",style='ba')
        x = i.split()
        if len(x) == 1:
            run = paragraph.add_run(i+":")
        else:
            run = paragraph.add_run(x[0]+'"'+x[1]+'"'+":")
            
        paragraph = document.add_paragraph("")
        run = paragraph.add_run("Descripción de la Unidad inmobiliaria: ")
        run.bold = True
        run.underline = True
        run = paragraph.add_run(dicdeff[i])
        document.add_paragraph("")
        
def add_Desc_Rep2(document,listofdatt,dicdeff, tipomod):
    for i in dicdeff:
        paragraph = document.add_paragraph("UNIDAD INMOBILIARIA ",style='ba')
        x = i.split()
        if len(x) == 1:
            run = paragraph.add_run(i+":")
        else:
            run = paragraph.add_run(x[0]+'"'+x[1]+'"'+":")

        paragraph = document.add_paragraph("")
        run = paragraph.add_run("Cantidad de Unidades Inmobiliarias y Nomenclatura: ")
        run.bold = True
        run.underline = True
        run = paragraph.add_run("Comprende de ")
        run = paragraph.add_run(get_numbandword( str(len(tipomod[i])  )))
        run = paragraph.add_run(" unidades inmobiliarias identificadas como:  ")
        ccc = 0
        for p in tipomod[i]:
            if ccc != len(tipomod[i])-1:
                run = paragraph.add_run(p+", ")
            else:
                run = paragraph.add_run(p+".")
            ccc += 1
        

        paragraph = document.add_paragraph("")
        paragraph = document.add_paragraph("")
        run = paragraph.add_run("Descripción de la Unidad inmobiliaria: ")
        run.bold = True
        run.underline = True
        run = paragraph.add_run(dicdeff[i])
        document.add_paragraph("")




def add_Desc_Etapa_Current(document,listofdatt,dict_letters,list_of_phase,totcount_phase):
    paragraph = document.add_paragraph("")
    run = paragraph.add_run("La construcción de la etapa numero ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][1])))
    run = paragraph.add_run(" del proyecto ocupará una superficie de ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][11])))
    run = paragraph.add_run(", de los cuales ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][12])))
    run = paragraph.add_run(" corresponden al área de construcción de las unidades inmobiliarias y ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][13])))
    run = paragraph.add_run("corresponden a áreas comunes que incluyen calles, gramas y aceras del proyecto.")
    paragraph = document.add_paragraph("")
    paragraph = document.add_paragraph("")
    run = paragraph.add_run("La construcción de la etapa numero ")
    run = paragraph.add_run(get_numbandword(str(listofdatt[0][1])))
    run = paragraph.add_run(" del proyecto estará compuesta por ")
    run = paragraph.add_run(get_numbandword(str(totcount_phase)))
    run = paragraph.add_run(" unidades inmobiliarias, las cuales serán identificadas como ")
    cccc = 0
    list_of_phase
    for p in list_of_phase:
        for i in list_of_phase[p]:
            i = list(i)
            run = paragraph.add_run(p+"-"+str(i[0])+" a la ")
            run = paragraph.add_run(p+"-"+str(i[1])+",")
            if cccc != countofPairs-1:
                run = paragraph.add_run(" y ")
            cccc += 1
    run = paragraph.add_run(" compuestas por ")
    
    for i in dict_letters:
        run = paragraph.add_run("Unidades Inmobiliaria Tipo ")
        run = paragraph.add_run(i+" ")
        for l in dict_letters[i]:
            run = paragraph.add_run('"'+l+'", ')
    run = paragraph.add_run(" los cuales se detallan a continuación:")
    
        
        
add_Title(document,listofdatt)
paragrap = document.add_paragraph(str())
add_Art_1(document,listofdatt)
paragrap = document.add_paragraph(str())
add_Art_2(document,listofdatt)
paragrap = document.add_paragraph(str())
add_Desc_Gen(document,listofdatt,dicTOT)
paragrap = document.add_paragraph(str())
add_Desc_Rep(document,listofdatt,dicdeff)
paragrap = document.add_paragraph(str())
add_Desc_Etapa_Current(document,listofdatt,dict_letters,list_of_phase,totcount_phase)
paragrap = document.add_paragraph(str())
add_Desc_Rep2(document,listofdatt,dicdeff, tipomod)
document.save("test.docx")


#
#print(listofdatt[0])
